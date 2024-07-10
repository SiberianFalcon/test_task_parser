import re
import os

from bs4 import BeautifulSoup
from docx import Document

from constants import UNICODE_SUBTITLE_NUMS, TEMPLATES, UNNECESSARY_TEXT, GARBAGE_LIST

current_dir = os.getcwd()
for set_of_templates in TEMPLATES:
    file_path = os.path.join(current_dir, set_of_templates)
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

    soup = BeautifulSoup(content, "lxml")

    tag_with_need_part = soup.find('div', {'id': 'premain'}).find('b')

    values = []
    fix_state = tag_with_need_part.next_element
    current_state = fix_state
    while True:
        get_next_tag = current_state.next_element
        if get_next_tag in UNNECESSARY_TEXT:
            break
        values.append(get_next_tag)
        current_state = get_next_tag

    get_all_expr = ''
    for i in range(len(values)):
        if i == len(values):
            break

        if re.match('Для построения', str(values[i])):
            break

        if re.match('Определим максимальное значение целевой функции ', str(values[i])):
            pars_string = re.split(
                'Определим максимальное значение целевой функции ', str(values[i])
            )
            get_all_expr += pars_string[1]

        if re.match('<sub>(?P<substring>\d+)</sub>', str(values[i])):
            get_subtitle = re.match('<sub>(?P<substring>\d+)</sub>', str(values[i]))
            change_subt_num = re.sub(
                '<sub>\d+</sub>',
                UNICODE_SUBTITLE_NUMS[get_subtitle.group("substring")],
                str(values[i]),
            )
            get_all_expr += change_subt_num

            if values[i+2] != ' при следующих условиях-ограничений.':
                get_all_expr += f'{values[i+2]}'
            del values[i+1]

        if str(values[i]) == '<br/>':
            if str(values[i+1]) not in GARBAGE_LIST:
                get_all_expr += f' {str(values[i+1])}'

    first_part = get_all_expr.split(maxsplit=3)[0:3]
    get_x = ''
    for i in range(len(first_part[2]) - 1):
        if first_part[2][i] == 'x':
            get_x += first_part[2][i]
            get_x += first_part[2][i+1] + ','

    replace_x = ''.join(first_part)
    result = replace_x.replace('X', get_x[:-1])
    result += ' \u2192 𝑚𝑎𝑥'
    others = get_all_expr.split()[3:]

    document = Document()
    document.add_paragraph(result)
    for i in range(len(others)):
        paragraph = document.add_paragraph()
        paragraph.text = others[i]

    document.save(f'parse_{set_of_templates[0:-5]}.docx')
